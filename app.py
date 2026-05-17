import streamlit as st
import docx
import re
import io

# ==========================================
# ANALYZER CLASS
# ==========================================

class ProjectFileAnalyzer:
    def __init__(self, file_bytes):
        """Initializes the analyzer and loads the Word document."""
        try:
            self.doc = docx.Document(io.BytesIO(file_bytes))
            self.full_text = [p.text for p in self.doc.paragraphs if p.text.strip()]
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in self.full_text:
                            self.full_text.append(text)
        except Exception as e:
            st.error(f"Error loading document: {e}")
            self.doc = None

    def check_front_pages(self, mentor_name, subject_code):
        """
        Scans front matter (up to first Heading 1 or 150 blocks) for:
          - Subject code, mentor name, institution, university,
            proforma, certificate, student roll number(s), glossary.
        """
        if not self.doc:
            return {}

        front_blocks = []
        for p in self.doc.paragraphs:
            if p.style and p.style.name == "Heading 1" and front_blocks:
                break
            front_blocks.append(p.text)
            if len(front_blocks) >= 150:
                break

        # Also scan table text in case details sit inside tables
        table_text_parts = []
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text_parts.append(cell.text.strip())

        front_text = " ".join(front_blocks + table_text_parts).lower()

        checks = {
            f"Subject Code ({subject_code})": False,
            f"Mentor / Guide Name ({mentor_name})": False,
            "Institution (IMS Noida)": False,
            "University (Chaudhary Charan Singh University)": False,
            "Proforma for Approval": False,
            "Certificate of Originality": False,
            "Student Roll Number(s) present": False,
            "Glossary section present": False,
        }

        if subject_code.lower() in front_text:
            checks[f"Subject Code ({subject_code})"] = True

        if mentor_name.lower() in front_text:
            checks[f"Mentor / Guide Name ({mentor_name})"] = True

        if "ims noida" in front_text or "institute of management studies" in front_text:
            checks["Institution (IMS Noida)"] = True

        if "chaudhary charan singh" in front_text or "ccsu" in front_text:
            checks["University (Chaudhary Charan Singh University)"] = True

        if "proforma" in front_text:
            checks["Proforma for Approval"] = True

        if "certificate of originality" in front_text:
            checks["Certificate of Originality"] = True

        # Roll numbers: CCSU format e.g. R220919106080
        if re.search(r'r\d{12}', front_text) or re.search(r'roll\s*(no|num|number)', front_text):
            checks["Student Roll Number(s) present"] = True

        # Glossary: anywhere in the full document
        full_doc_text = " ".join(self.full_text).lower()
        if "glossary" in full_doc_text:
            checks["Glossary section present"] = True

        return checks

    def verify_index_table(self):
        """
        Verifies the Index table has:
          - 3 columns: Sr. No. | Description | Page No.
          - All 11 mandatory BCA major project chapters (per official template).
        """
        if not self.doc:
            return False, "Document not loaded."

        index_table = None

        for table in self.doc.tables:
            if len(table.rows) > 0:
                first_row_text = [cell.text.strip().lower() for cell in table.rows[0].cells]
                has_description = any("description" in t for t in first_row_text)
                has_page = any("page" in t for t in first_row_text)
                if has_description and has_page:
                    index_table = table
                    break

        if not index_table:
            return (
                False,
                "No structured Index table found. The Index must be a table with columns: "
                "'Sr. No.', 'Description', and 'Page No.'",
            )

        # Check column count (should be 3)
        col_count = len(index_table.columns)
        col_warning = ""
        if col_count != 3:
            col_warning = (
                f" ⚠️ Index table has {col_count} column(s); "
                "expected 3 (Sr. No. | Description | Page No.)."
            )

        # Extract all cell text
        table_text = []
        for row in index_table.rows:
            for cell in row.cells:
                table_text.append(cell.text.strip().lower())
        full_index_text = " ".join(table_text)

        # Official BCA-605P mandatory chapters (from actual PDF template)
        benchmark_chapters = [
            "about project",
            "system analysis",
            "system design",
            "coding",
            "testing",
            "system security",
            "cost estimation",
            "reports",
            "future scope",
            "limitations",
            "bibliography",
        ]

        missing_chapters = [
            chapter.title()
            for chapter in benchmark_chapters
            if chapter not in full_index_text
        ]

        if missing_chapters:
            missing_str = ", ".join(missing_chapters)
            return (
                False,
                f"Index table found, but missing mandatory chapters: {missing_str}.{col_warning}",
            )

        return (
            True,
            f"Index table matches the standard format — all 11 mandatory chapters present!{col_warning}",
        )

    def check_mandatory_sections(self):
        """
        Checks the document body for the actual BCA major project sections
        (matching the official PDF template — NOT generic research paper sections).
        Only matches against heading-style paragraphs or short standalone lines.
        """
        if not self.doc:
            return {}

        # Sections as they appear in the actual BCA-605P template
        required_sections = [
            "Introduction",
            "System Analysis",
            "System Design",
            "Coding",
            "Testing",
            "System Security",
            "Cost Estimation",
            "Reports",
            "Future Scope",
            "Limitations",
            "Bibliography",
        ]
        found_sections = {section: False for section in required_sections}

        for paragraph in self.doc.paragraphs:
            is_heading = bool(paragraph.style and paragraph.style.name.startswith("Heading"))
            text = paragraph.text.strip().lower()
            for section in required_sections:
                if section.lower() in text:
                    # Accept headings, or short standalone lines (≤ 60 chars)
                    if is_heading or len(text) <= 60:
                        found_sections[section] = True

        return found_sections

    def analyze_structure_and_headings(self):
        """Extracts paragraphs using Word's built-in Heading styles."""
        if not self.doc:
            return []
        headings = []
        for paragraph in self.doc.paragraphs:
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                headings.append(
                    {"level": paragraph.style.name, "text": paragraph.text.strip()}
                )
        return headings

    def check_introduction_length(self, min_words=200, max_words=800):
        """
        Checks the Introduction section word count.
        (Replaces the generic 'Abstract' check — BCA template starts with Introduction.)
        """
        if not self.doc:
            return None

        intro_text = ""
        is_intro = False

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text.lower().startswith("introduction") and not is_intro:
                is_intro = True
                continue
            if is_intro and paragraph.style and paragraph.style.name.startswith("Heading"):
                break
            if is_intro:
                intro_text += text + " "

        word_count = len(intro_text.split())
        if word_count == 0:
            return "Introduction section not found or empty.", "error"
        elif min_words <= word_count <= max_words:
            return f"Pass ({word_count} words — within {min_words}–{max_words} range)", "success"
        else:
            return (
                f"Fail ({word_count} words). Recommended range: {min_words}–{max_words} words.",
                "warning",
            )

    def check_bibliography(self):
        """Checks that Bibliography section exists and has at least 5 entries."""
        if not self.doc:
            return None

        bib_text = ""
        is_bib = False

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text.lower().startswith("bibliography") or text.lower().startswith("references"):
                is_bib = True
                continue
            if is_bib:
                bib_text += text + "\n"

        if not bib_text.strip():
            return "Bibliography / References section not found or empty.", "error"

        entries = [
            line.strip()
            for line in bib_text.splitlines()
            if line.strip() and len(line.strip()) > 10
        ]
        count = len(entries)

        if count >= 5:
            return f"Pass — {count} reference entries found.", "success"
        else:
            return (
                f"Warn — only {count} reference entries found. Aim for at least 5.",
                "warning",
            )


# ==========================================
# SUBJECT CODE OPTIONS
# ==========================================

SUBJECT_CODES = [
    "BCA-605",
    "Custom — enter below",
]

# ==========================================
# STREAMLIT USER INTERFACE
# ==========================================

st.set_page_config(
    page_title="Project Report Analyzer", page_icon="📄", layout="centered"
)

st.title("📄 Major Project Report Analyzer")
st.write(
    "Ensure your project file meets the formatting, structural, and institutional "
    "guidelines before submission."
)

# --- Student Input Fields ---
st.info("Please provide your specific project details before uploading your file.")

col_input1, col_input2 = st.columns(2)

with col_input1:
    mentor_name = st.text_input(
        "Mentor's / Guide's Name", placeholder="e.g., Mr. Prateek Tiwari"
    )

with col_input2:
    selected_code = st.selectbox("Course & Subject Code", SUBJECT_CODES)
    if selected_code == "Custom — enter below":
        subject_code = st.text_input(
            "Enter custom subject code", placeholder="e.g., BBA-705P"
        )
    else:
        subject_code = selected_code
        st.caption(f"Selected: **{subject_code}**")

uploaded_file = st.file_uploader(
    "Upload your Word Document (.docx)", type=["docx"]
)

if uploaded_file is not None:
    if not mentor_name or not subject_code:
        st.warning(
            "⚠️ Please enter both your Mentor's Name and Subject Code above to begin the analysis."
        )
    else:
        file_bytes = uploaded_file.read()

        with st.spinner("Analyzing document structure and content..."):
            analyzer = ProjectFileAnalyzer(file_bytes)

            if analyzer.doc:
                st.divider()

                # ── 1. Title Page & Certificate Checks ──────────────────────
                st.header("1. Title Page & Certificate Checks")
                st.write(
                    "Checks: Subject Code · Mentor Name · Institution · University · "
                    "Proforma · Certificate of Originality · Roll Numbers · Glossary"
                )
                front_page_results = analyzer.check_front_pages(mentor_name, subject_code)

                passed = sum(1 for v in front_page_results.values() if v)
                total = len(front_page_results)
                st.caption(f"{passed}/{total} checks passed")

                for detail, is_present in front_page_results.items():
                    if is_present:
                        st.success(f"✅ Found: {detail}")
                    else:
                        st.error(f"❌ Missing: {detail}")

                st.divider()

                # ── 2. Index Verification ────────────────────────────────────
                st.header("2. Index Verification")
                st.write(
                    "Checks that the Index is a proper 3-column table "
                    "(Sr. No. | Description | Page No.) with all 11 mandatory chapters."
                )

                index_passed, index_msg = analyzer.verify_index_table()

                if index_passed:
                    st.success(f"✅ {index_msg}")
                    st.info(
                        "💡 Verify page numbers manually — Word renders them dynamically "
                        "and they cannot be auto-checked."
                    )
                else:
                    st.error(f"❌ {index_msg}")
                    with st.expander("📋 View Required Index Format"):
                        st.markdown(
                            """
The Index must be a **3-column table** with headers:
`Sr. No.` | `Description` | `Page No.`

It must include all of these chapters:

| # | Chapter |
|---|---------|
| 1 | About Project |
| 2 | System Analysis |
| 3 | System Design |
| 4 | Coding |
| 5 | Testing |
| 6 | System Security Measures |
| 7 | Cost Estimation of the Project |
| 8 | Reports |
| 9 | Future Scope and Further Enhancement |
| 10 | Limitations of the Project |
| 11 | Bibliography |
                            """
                        )

                st.divider()

                # ── 3. General Word Count ────────────────────────────────────
                st.header("3. General Word Count")
                total_words = sum(len(text.split()) for text in analyzer.full_text)
                if total_words >= 8000:
                    st.success(f"✅ Total Word Count: **{total_words:,} words** — meets minimum length.")
                elif total_words >= 5000:
                    st.warning(f"⚠️ Total Word Count: **{total_words:,} words** — consider expanding content.")
                else:
                    st.error(f"❌ Total Word Count: **{total_words:,} words** — document appears too short.")

                st.divider()

                # ── 4. Mandatory BCA Sections ────────────────────────────────
                st.header("4. Mandatory Project Sections")
                st.write(
                    "Checks for the 11 official BCA major project sections "
                    "as per the IMS Noida / CCSU template."
                )
                sections = analyzer.check_mandatory_sections()

                passed_s = sum(1 for v in sections.values() if v)
                total_s = len(sections)
                st.caption(f"{passed_s}/{total_s} sections found")

                col1, col2 = st.columns(2)
                with col1:
                    st.write("**Required Section**")
                with col2:
                    st.write("**Status**")

                for section, is_present in sections.items():
                    colA, colB = st.columns(2)
                    colA.write(section)
                    if is_present:
                        colB.success("✅ Found")
                    else:
                        colB.error("❌ Missing")

                st.divider()

                # ── 5. Content Rules ─────────────────────────────────────────
                st.header("5. Content Rules")

                intro_result = analyzer.check_introduction_length()
                if intro_result:
                    msg, status = intro_result
                    label = "**Introduction Length:**"
                    if status == "success":
                        st.success(f"✅ {label} {msg}")
                    elif status == "warning":
                        st.warning(f"⚠️ {label} {msg}")
                    else:
                        st.error(f"❌ {label} {msg}")

                bib_result = analyzer.check_bibliography()
                if bib_result:
                    msg, status = bib_result
                    label = "**Bibliography / References:**"
                    if status == "success":
                        st.success(f"✅ {label} {msg}")
                    elif status == "warning":
                        st.warning(f"⚠️ {label} {msg}")
                    else:
                        st.error(f"❌ {label} {msg}")

                st.divider()

                # ── 6. Document Structure (Headings) ─────────────────────────
                st.header("6. Document Structure (Headings)")
                st.write(
                    "Sections listed here used Word's official Heading styles. "
                    "Missing sections likely used bold text instead — fix those in Word."
                )
                headings = analyzer.analyze_structure_and_headings()

                if not headings:
                    st.error(
                        "❌ No formal headings detected. "
                        "Apply Word's Heading 1 / Heading 2 styles instead of manually bolding text."
                    )
                else:
                    st.success(f"✅ {len(headings)} heading(s) detected using proper Word styles.")
                    with st.expander("📑 View All Extracted Headings"):
                        for h in headings:
                            level_match = re.search(r"\d+", h["level"])
                            indent_level = int(level_match.group()) if level_match else 1
                            indent = "&nbsp;" * (indent_level * 4)
                            st.markdown(
                                f"{indent}▸ **[{h['level']}]** {h['text']}",
                                unsafe_allow_html=True,
                            )

# ==========================================
# FOOTER
# ==========================================

st.divider()
st.markdown(
    """
    <div style="text-align: center; color: #888; font-size: 0.85rem; padding: 8px 0 4px;">
        Built by&nbsp;
        <a href="https://www.linkedin.com/in/anushrav-mudgal/" target="_blank"
           style="color: #0A66C2; text-decoration: none; font-weight: 600;">
            Anushrav Mudgal
        </a>
        &nbsp;·&nbsp;
        <a href="https://www.linkedin.com/in/anushrav-mudgal/" target="_blank"
           style="color: #0A66C2; text-decoration: none;">
            🔗 LinkedIn
        </a>
    </div>
    """,
    unsafe_allow_html=True,
)
